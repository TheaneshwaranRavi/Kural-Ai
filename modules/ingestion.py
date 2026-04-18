import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {".pdf", ".txt", ".docx"}

MATH_SYMBOL_MAP = {
    "±": "plus or minus",
    "≠": "not equal to",
    "≤": "less than or equal to",
    "≥": "greater than or equal to",
    "≈": "approximately equal to",
    "∑": "sum of",
    "∏": "product of",
    "√": "square root of",
    "∞": "infinity",
    "π": "pi",
    "α": "alpha",
    "β": "beta",
    "γ": "gamma",
    "δ": "delta",
    "θ": "theta",
    "μ": "mu",
    "σ": "sigma",
    "∈": "element of",
    "∉": "not element of",
    "∩": "intersection",
    "∪": "union",
    "∀": "for all",
    "∃": "there exists",
    "→": "implies",
    "↔": "if and only if",
    "∧": "and",
    "∨": "or",
    "¬": "not",
    "²": "squared",
    "³": "cubed",
    "°": "degrees",
    "%": "percent",
    "÷": "divided by",
    "×": "multiplied by",
    "·": "multiplied by",
}


@dataclass
class ProcessedDocument:
    file_path: str
    exam_type: str
    subject: str
    topic: str
    language: str
    raw_text: str
    pages: List[str] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    ocr_used: bool = False
    page_count: int = 0
    char_count: int = 0


class DocumentProcessor:
    def __init__(self, ocr_languages: str = "eng+tam"):
        self._ocr_languages = ocr_languages
        self._tesseract_available = self._check_tesseract()

    def _check_tesseract(self) -> bool:
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            logger.warning(
                "Tesseract not found. OCR disabled. "
                "Install with: brew install tesseract tesseract-lang (macOS) "
                "or: sudo apt install tesseract-ocr tesseract-ocr-tam (Ubuntu)"
            )
            return False

    def process_document(
        self,
        file_path: str,
        exam_type: str,
        subject: str,
        topic: str,
        language: str = "en",
    ) -> ProcessedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format '{suffix}'. Supported: {SUPPORTED_FORMATS}"
            )

        logger.info(f"Processing {path.name} [{exam_type} | {subject} | {topic}]")

        doc = ProcessedDocument(
            file_path=str(path),
            exam_type=exam_type,
            subject=subject,
            topic=topic,
            language=language,
            raw_text="",
        )

        if suffix == ".pdf":
            self._load_pdf(path, doc)
        elif suffix == ".txt":
            self._load_txt(path, doc)
        elif suffix == ".docx":
            self._load_docx(path, doc)

        doc.raw_text = self._clean_text(doc.raw_text)
        doc.char_count = len(doc.raw_text)
        logger.info(
            f"Processed {path.name}: {doc.page_count} pages, "
            f"{doc.char_count} chars, OCR={doc.ocr_used}"
        )
        return doc

    def _load_pdf(self, path: Path, doc: ProcessedDocument) -> None:
        import fitz

        pdf = fitz.open(str(path))
        doc.page_count = len(pdf)
        all_pages: List[str] = []

        for page_num, page in enumerate(pdf, start=1):
            page_text = page.get_text("text").strip()

            if len(page_text) < 50 and self._tesseract_available:
                logger.debug(f"Page {page_num}: sparse text, falling back to OCR")
                page_text = self._ocr_pdf_page(page)
                doc.ocr_used = True

            tables_on_page = self._extract_tables_pdf(page)
            for table_str in tables_on_page:
                doc.tables.append(table_str)
                page_text += f"\n\n{table_str}"

            page_text = self._expand_math_symbols(page_text)
            all_pages.append(page_text)

        pdf.close()
        doc.pages = all_pages
        doc.raw_text = "\n\n".join(all_pages)

    def _ocr_pdf_page(self, page) -> str:
        import pytesseract
        from PIL import Image

        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_bytes))
        text = pytesseract.image_to_string(image, lang=self._ocr_languages)
        return text.strip()

    def _extract_tables_pdf(self, page) -> List[str]:
        tables: List[str] = []
        try:
            table_finder = page.find_tables()
            for table in table_finder.tables:
                rows = table.extract()
                if rows:
                    tables.append(self._format_table(rows))
        except Exception as e:
            logger.debug(f"Table extraction skipped: {e}")
        return tables

    def _load_txt(self, path: Path, doc: ProcessedDocument) -> None:
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = path.read_bytes().decode("utf-8", errors="replace")

        doc.raw_text = self._expand_math_symbols(text)
        doc.page_count = 1
        doc.pages = [doc.raw_text]

    def _load_docx(self, path: Path, doc: ProcessedDocument) -> None:
        from docx import Document as DocxDocument

        docx = DocxDocument(str(path))
        sections: List[str] = []

        for element in docx.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                from docx.oxml.ns import qn
                para_text = "".join(
                    node.text or ""
                    for node in element.iter()
                    if node.tag == qn("w:t")
                )
                if para_text.strip():
                    sections.append(self._expand_math_symbols(para_text.strip()))

            elif tag == "tbl":
                rows = []
                for row_el in element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                    cells = []
                    for cell_el in row_el.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc"):
                        cell_text = "".join(
                            t.text or ""
                            for t in cell_el.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                        ).strip()
                        cells.append(cell_text)
                    if cells:
                        rows.append(cells)
                if rows:
                    table_str = self._format_table(rows)
                    doc.tables.append(table_str)
                    sections.append(table_str)

        doc.raw_text = "\n\n".join(sections)
        doc.page_count = 1
        doc.pages = [doc.raw_text]

    def _format_table(self, rows: List[List]) -> str:
        if not rows:
            return ""

        str_rows = [[str(cell) if cell is not None else "" for cell in row] for row in rows]
        col_count = max(len(r) for r in str_rows)
        str_rows = [r + [""] * (col_count - len(r)) for r in str_rows]

        col_widths = [
            max(len(str_rows[r][c]) for r in range(len(str_rows)))
            for c in range(col_count)
        ]

        lines: List[str] = []
        for i, row in enumerate(str_rows):
            cells = [row[c].ljust(col_widths[c]) for c in range(col_count)]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("|" + "|".join("-" * (w + 2) for w in col_widths) + "|")

        return "\n".join(lines)

    def _expand_math_symbols(self, text: str) -> str:
        for symbol, description in MATH_SYMBOL_MAP.items():
            text = text.replace(symbol, f" {description} ")
        return text

    def _clean_text(self, text: str) -> str:
        text = unicodedata.normalize("NFC", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\S\n]+\n", "\n", text)
        text = re.sub(r"\n[^\S\n]+", "\n", text)
        text = re.sub(r"\.{4,}", "...", text)
        return text.strip()


def process_document(
    file_path: str,
    exam_type: str,
    subject: str,
    topic: str,
    language: str = "en",
    ocr_languages: str = "eng+tam",
) -> ProcessedDocument:
    processor = DocumentProcessor(ocr_languages=ocr_languages)
    return processor.process_document(file_path, exam_type, subject, topic, language)
